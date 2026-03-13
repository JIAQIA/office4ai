"""
Word E2E 测试 Fixture 生成脚本

程序化生成所有 Word E2E 测试所需的 fixture .docx 文件。

Usage:
    uv run python manual_tests/word/fixtures/create_fixtures.py
    uv run python manual_tests/word/fixtures/create_fixtures.py --clean  # 清理后重新生成
"""

from pathlib import Path

from docx import Document

FIXTURES_ROOT = Path(__file__).parent


# ==============================================================================
# insert_text_e2e fixtures
# ==============================================================================


def create_insert_text_fixtures() -> None:
    """创建 insert_text_e2e 所需的 fixture 文件"""
    fixture_dir = FIXTURES_ROOT / "insert_text_e2e"
    fixture_dir.mkdir(parents=True, exist_ok=True)

    # empty.docx — 空文档
    doc = Document()
    doc.save(str(fixture_dir / "empty.docx"))

    # simple.docx — 3 段已知文本
    doc = Document()
    doc.add_paragraph("这是第一段测试文本。")
    doc.add_paragraph("这是第二段测试文本。")
    doc.add_paragraph("这是第三段测试文本。")
    doc.save(str(fixture_dir / "simple.docx"))

    print("  ✅ insert_text_e2e: 2 files")


# ==============================================================================
# replace_text_e2e fixtures
# ==============================================================================


def create_replace_text_fixtures() -> None:
    """创建 replace_text_e2e 所需的 fixture 文件"""
    fixture_dir = FIXTURES_ROOT / "replace_text_e2e"
    fixture_dir.mkdir(parents=True, exist_ok=True)

    # replace_targets.docx — 包含多种可搜索文本
    doc = Document()
    doc.add_heading("Replace Test Document", level=1)

    # "old" x5
    for i in range(5):
        doc.add_paragraph(f"This is old text number {i + 1}.")

    # "test" x5
    for i in range(5):
        doc.add_paragraph(f"This is a test sentence number {i + 1}.")

    # "delete" x3
    for i in range(3):
        doc.add_paragraph(f"Please delete this text {i + 1}.")

    # "Café" x3
    for i in range(3):
        doc.add_paragraph(f"Visit our Café for coffee {i + 1}.")

    # "line1" + "line2" 跨段落（用于 ^p 跨段落搜索测试）
    doc.add_paragraph("line1")
    doc.add_paragraph("line2")

    # 长段落
    long_text = (
        "This is a long paragraph of text that should be replaced with another "
        "long paragraph. It contains multiple sentences and various punctuation marks."
    )
    doc.add_paragraph(long_text)
    doc.add_paragraph(long_text)

    doc.save(str(fixture_dir / "replace_targets.docx"))

    # format_targets.docx — 用于格式化替换测试
    doc = Document()
    doc.add_heading("Format Replace Test Document", level=1)

    # 用于测试粗体格式化
    for i in range(3):
        doc.add_paragraph(f"Mark this important text as bold ({i + 1}).")

    # 用于测试斜体格式化
    for i in range(3):
        doc.add_paragraph(f"This emphasis text should be italic ({i + 1}).")

    # 用于测试颜色格式化
    for i in range(3):
        doc.add_paragraph(f"Warning: this alert text needs color ({i + 1}).")

    # 用于测试 styleName 格式化
    doc.add_paragraph("Chapter One Introduction")
    doc.add_paragraph("Chapter Two Methods")

    # 用于测试组合格式
    for i in range(2):
        doc.add_paragraph(f"Critical notice requiring combined formatting ({i + 1}).")

    doc.save(str(fixture_dir / "format_targets.docx"))

    print("  ✅ replace_text_e2e: 2 files")


# ==============================================================================
# replace_selection_e2e fixtures
# ==============================================================================


def create_replace_selection_fixtures() -> None:
    """创建 replace_selection_e2e 所需的 fixture 文件"""
    fixture_dir = FIXTURES_ROOT / "replace_selection_e2e"
    fixture_dir.mkdir(parents=True, exist_ok=True)

    # simple.docx — 4 段已知文本，供选择替换
    doc = Document()
    doc.add_paragraph("第一段：这是用于选择替换测试的文本。")
    doc.add_paragraph("第二段：请选中这段文字进行替换。")
    doc.add_paragraph("第三段：替换测试需要先选中文本。")
    doc.add_paragraph("第四段：最后一段测试文本。")
    doc.save(str(fixture_dir / "simple.docx"))

    # empty.docx — 空文档（edge case）
    doc = Document()
    doc.save(str(fixture_dir / "empty.docx"))

    print("  ✅ replace_selection_e2e: 2 files")


# ==============================================================================
# get_visible_content_e2e fixtures
# ==============================================================================


def create_get_visible_content_fixtures() -> None:
    """创建 get_visible_content_e2e 所需的 fixture 文件"""
    fixture_dir = FIXTURES_ROOT / "get_visible_content_e2e"
    fixture_dir.mkdir(parents=True, exist_ok=True)

    # empty.docx
    doc = Document()
    doc.save(str(fixture_dir / "empty.docx"))

    # simple.docx — 纯文本 3 段
    doc = Document()
    doc.add_paragraph("这是第一段简单文本。用于测试可见内容获取。")
    doc.add_paragraph("第二段文本在这里。包含一些中文和English混合内容。")
    doc.add_paragraph("最后一段文本。测试结束。")
    doc.save(str(fixture_dir / "simple.docx"))

    # complex.docx — 文本 + 表格 + 格式化
    doc = Document()
    doc.add_heading("复杂测试文档", level=1)
    doc.add_paragraph("这是一个包含多种元素的复杂文档。")
    doc.add_heading("表格示例", level=2)

    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"单元格 {i + 1}-{j + 1}"

    doc.add_heading("格式化文本", level=2)
    p = doc.add_paragraph()
    run = p.add_run("粗体文本")
    run.bold = True
    p.add_run(" 和 ")
    run = p.add_run("斜体文本")
    run.italic = True

    doc.add_paragraph("这是文档的结尾。")
    doc.save(str(fixture_dir / "complex.docx"))

    # large.docx — 10 页重复内容
    doc = Document()
    doc.add_heading("大型测试文档", level=1)
    para_text = (
        "这是一段用于填充大型文档的文本。"
        "我们需要足够多的内容来测试文档获取功能的性能。"
        "Word 文档可能包含大量的文字、段落和其他元素。"
        "这个测试旨在验证系统在处理大型文档时的表现。"
    ) * 5
    for i in range(30):  # ~10 pages
        doc.add_paragraph(f"第 {i + 1} 段：{para_text}")
    doc.save(str(fixture_dir / "large.docx"))

    print("  ✅ get_visible_content_e2e: 4 files")


# ==============================================================================
# select_text_e2e fixtures
# ==============================================================================


def create_select_text_fixtures() -> None:
    """创建 select_text_e2e 所需的 fixture 文件"""
    fixture_dir = FIXTURES_ROOT / "select_text_e2e"
    fixture_dir.mkdir(parents=True, exist_ok=True)

    # simple.docx — "Hello World"×3, "test"×5, 大小写变体
    doc = Document()
    doc.add_paragraph("Hello World is a common greeting.")
    doc.add_paragraph("Another Hello World appears here.")
    doc.add_paragraph("Third Hello World for testing.")
    doc.add_paragraph("This is a test sentence.")
    doc.add_paragraph("Another test for validation.")
    doc.add_paragraph("test number three here.")
    doc.add_paragraph("More test data follows.")
    doc.add_paragraph("Final test in this document.")
    doc.add_paragraph("Hello in lowercase.")
    doc.add_paragraph("HELLO in uppercase.")
    doc.add_paragraph("hello in all lowercase.")
    doc.add_paragraph("Selection Test for mode testing.")
    doc.add_paragraph("CursorPosition for start mode.")
    doc.add_paragraph("EndPosition for end mode.")
    doc.add_paragraph("ModeSwitch for switching modes.")
    doc.add_paragraph("Pattern as a complete word.")
    doc.add_paragraph("test123 is not a whole word match.")
    doc.add_paragraph("mytest is also not a whole word match.")
    doc.save(str(fixture_dir / "simple.docx"))

    # edge_cases.docx — 特殊字符、长文本
    doc = Document()
    doc.add_paragraph("Special chars: @#$%")
    doc.add_paragraph("Email: test@example.com")
    doc.add_paragraph("Path: C:\\Users\\test")
    doc.add_paragraph("Parenthesis: (parenthesis)")
    doc.add_paragraph("Brackets: [brackets]")
    doc.add_paragraph("Braces: {braces}")
    long_text = "This is a very long text that " * 10
    doc.add_paragraph(long_text)
    doc.add_paragraph("OutOfBounds text appears here.")
    doc.add_paragraph("OutOfBounds again for testing.")
    doc.add_paragraph("OutOfBounds third occurrence.")
    doc.save(str(fixture_dir / "edge_cases.docx"))

    print("  ✅ select_text_e2e: 2 files")


# ==============================================================================
# export_content_e2e fixtures
# ==============================================================================


def create_export_content_fixtures() -> None:
    """创建 export_content_e2e 所需的 fixture 文件

    复用 get_visible_content_e2e 的文档结构: empty, simple, complex, large。
    """
    fixture_dir = FIXTURES_ROOT / "export_content_e2e"
    fixture_dir.mkdir(parents=True, exist_ok=True)

    # empty.docx
    doc = Document()
    doc.save(str(fixture_dir / "empty.docx"))

    # simple.docx — 纯文本 3 段
    doc = Document()
    doc.add_paragraph("这是第一段简单文本。用于测试文档导出功能。")
    doc.add_paragraph("第二段文本在这里。包含一些中文和English混合内容。")
    doc.add_paragraph("最后一段文本。测试结束。")
    doc.save(str(fixture_dir / "simple.docx"))

    # complex.docx — 文本 + 表格 + 格式化
    doc = Document()
    doc.add_heading("复杂测试文档", level=1)
    doc.add_paragraph("这是一个包含多种元素的复杂文档。")
    doc.add_heading("表格示例", level=2)

    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"单元格 {i + 1}-{j + 1}"

    doc.add_heading("格式化文本", level=2)
    p = doc.add_paragraph()
    run = p.add_run("粗体文本")
    run.bold = True
    p.add_run(" 和 ")
    run = p.add_run("斜体文本")
    run.italic = True

    doc.add_paragraph("这是文档的结尾。")
    doc.save(str(fixture_dir / "complex.docx"))

    # large.docx — 10 页重复内容
    doc = Document()
    doc.add_heading("大型测试文档", level=1)
    para_text = (
        "这是一段用于填充大型文档的文本。"
        "我们需要足够多的内容来测试文档导出功能的性能。"
        "Word 文档可能包含大量的文字、段落和其他元素。"
        "这个测试旨在验证系统在处理大型文档时的表现。"
    ) * 5
    for i in range(30):  # ~10 pages
        doc.add_paragraph(f"第 {i + 1} 段：{para_text}")
    doc.save(str(fixture_dir / "large.docx"))

    print("  ✅ export_content_e2e: 4 files")


# ==============================================================================
# comment_e2e fixtures
# ==============================================================================


def create_comment_fixtures() -> None:
    """创建 comment_e2e 所需的 fixture 文件"""
    fixture_dir = FIXTURES_ROOT / "comment_e2e"
    fixture_dir.mkdir(parents=True, exist_ok=True)

    # simple.docx — 3 段已知文本（用于搜索定位批注）
    doc = Document()
    doc.add_paragraph("这是第一段测试文本。用于批注功能的端到端测试。")
    doc.add_paragraph("第二段文本包含一些关键词。测试批注定位功能。")
    doc.add_paragraph("最后一段文本。批注测试结束。")
    doc.save(str(fixture_dir / "simple.docx"))

    # empty.docx — 空文档
    doc = Document()
    doc.save(str(fixture_dir / "empty.docx"))

    print("  ✅ comment_e2e: 2 files")


# ==============================================================================
# Main
# ==============================================================================


def main() -> None:
    """生成所有 fixture 文件"""
    import argparse

    parser = argparse.ArgumentParser(description="生成 E2E 测试 fixture 文件")
    parser.add_argument("--clean", action="store_true", help="清理后重新生成")
    args = parser.parse_args()

    if args.clean:
        for subdir in FIXTURES_ROOT.iterdir():
            if subdir.is_dir() and subdir.name != "__pycache__":
                for docx in subdir.glob("*.docx"):
                    docx.unlink()
                    print(f"  🗑️  删除: {docx.relative_to(FIXTURES_ROOT)}")

    print("📝 生成 Word E2E 测试 fixture 文件...\n")

    create_insert_text_fixtures()
    create_replace_text_fixtures()
    create_replace_selection_fixtures()
    create_get_visible_content_fixtures()
    create_select_text_fixtures()
    create_export_content_fixtures()
    create_comment_fixtures()

    print("\n✅ Word fixture 全部完成")


if __name__ == "__main__":
    main()
