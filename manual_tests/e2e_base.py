"""
E2E 测试基础设施

提供自动化的文档夹具管理、自动打开、验证和清理功能。

使用方式:
    from manual_tests.e2e_base import E2ETestRunner, DocumentFixture

    runner = E2ETestRunner()
    async with runner.prepare_document("fixtures/empty.docx") as fixture:
        # fixture.document_uri 是复制后的文档 URI
        # 文档已经自动打开
        result = await run_test(fixture)
        # 成功后自动删除，失败则保留
"""

import asyncio
import inspect
import platform
import shutil
import subprocess
from collections.abc import AsyncIterator, Callable
from contextlib import asynccontextmanager
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING, Any
from urllib.parse import quote

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument

from office4ai.environment.workspace.office_workspace import OfficeWorkspace

# ==============================================================================
# 配置常量
# ==============================================================================

# Add-In 名称（用于自动激活）
DEFAULT_ADDIN_NAME = "word-editor"

# 测试夹具根目录
FIXTURES_ROOT = Path(__file__).parent / "fixtures"

# 临时文件目录（测试副本存放位置）
# 使用项目内的目录，避免系统临时目录的路径解析问题（如 macOS /var → /private/var）
# 同时方便调试时查看测试文件
TEMP_ROOT = Path(__file__).parent / ".test_working"


# ==============================================================================
# 数据类
# ==============================================================================


@dataclass
class DocumentReader:
    """
    文档内容读取器

    提供对测试文档内容的只读访问，用于双重验证。
    通过 python-docx 读取修改后的文档，验证实际内容。

    当 auto_save=True 时，reload() 会先通过 AppleScript 强制 Word 保存文档，
    确保磁盘文件与 Word 内存中的内容一致。

    Example:
        reader = DocumentReader(fixture.working_path)
        assert reader.contains("预期文本")
        assert reader.paragraph_contains(0, "第一段内容")
    """

    path: Path
    auto_save: bool = True
    _doc: "DocxDocument | None" = field(default=None, repr=False)

    @property
    def doc(self) -> "DocxDocument":
        """懒加载 Document 对象"""
        if self._doc is None:
            from docx import Document

            self._doc = Document(str(self.path))
        return self._doc

    def reload(self) -> None:
        """
        重新加载文档

        当 auto_save=True 时，先通过 AppleScript 强制 Word 保存文档到磁盘，
        再重新加载。这解决了 Add-In 修改文档后内存与磁盘不同步的问题。
        """
        if self.auto_save:
            save_document(self.path)
        self._doc = None

    @property
    def paragraphs(self) -> list[str]:
        """获取所有段落文本"""
        return [p.text for p in self.doc.paragraphs]

    @property
    def text(self) -> str:
        """获取全文文本（段落用换行连接）"""
        return "\n".join(self.paragraphs)

    @property
    def table_count(self) -> int:
        """获取表格数量"""
        return len(self.doc.tables)

    def contains(self, text: str) -> bool:
        """检查文档是否包含指定文本"""
        return text in self.text

    def paragraph_contains(self, index: int, text: str) -> bool:
        """
        检查指定段落是否包含文本

        Args:
            index: 段落索引（0-based）
            text: 要搜索的文本

        Returns:
            如果段落存在且包含文本返回 True，否则返回 False
        """
        if 0 <= index < len(self.doc.paragraphs):
            return text in self.doc.paragraphs[index].text
        return False

    def get_paragraph(self, index: int) -> str | None:
        """
        获取指定段落的文本

        Args:
            index: 段落索引（0-based）

        Returns:
            段落文本，如果索引无效返回 None
        """
        if 0 <= index < len(self.doc.paragraphs):
            return self.doc.paragraphs[index].text
        return None

    def not_contains(self, text: str) -> bool:
        """检查文档不包含指定文本（replace 验证用）"""
        return text not in self.text

    def paragraph_starts_with(self, index: int, text: str) -> bool:
        """
        检查段落是否以指定文本开头（location=Start 验证用）

        Args:
            index: 段落索引（0-based）
            text: 要检查的前缀文本

        Returns:
            如果段落存在且以指定文本开头返回 True，否则返回 False
        """
        if 0 <= index < len(self.doc.paragraphs):
            return self.doc.paragraphs[index].text.startswith(text)
        return False

    def paragraph_has_style(self, text: str, style_name: str) -> bool:
        """
        检查包含指定文本的段落是否使用了指定样式

        Args:
            text: 要搜索的文本
            style_name: 预期的段落样式名（如 "Heading 2", "Normal"）

        Returns:
            如果找到包含文本的段落且样式匹配返回 True
        """
        for paragraph in self.doc.paragraphs:
            if text in paragraph.text:
                if paragraph.style and paragraph.style.name == style_name:
                    return True
        return False

    def run_has_style(self, text: str, style_name: str) -> bool:
        """
        检查包含指定文本的 run 是否被应用了指定的字符样式

        Word 的 replace:text + format.styleName 会将段落样式的字符部分
        以 rStyle 应用到匹配的 run 上。例如 "Heading 2" 会产生
        "Heading 2 Char"（英文）或 "标题 2 字符"（中文）。

        本方法通过匹配 style_name 子串来兼容不同 locale。

        Args:
            text: 要搜索的文本
            style_name: 样式名关键词（如 "Heading 2" 或 "标题 2"）

        Returns:
            如果找到匹配的 run 且其 rStyle 包含关键词返回 True
        """
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                if text in run.text and run.style and run.style.name:
                    if style_name in run.style.name:
                        return True
        return False

    def run_has_format(
        self,
        text: str,
        *,
        bold: bool | None = None,
        italic: bool | None = None,
        font_name: str | None = None,
        font_size: float | None = None,
    ) -> bool:
        """
        检查包含指定文本的 run 是否具有指定格式

        在文档所有段落的所有 run 中搜索包含 text 的 run，
        验证其格式属性是否匹配。仅检查非 None 的参数。

        Args:
            text: 要搜索的文本
            bold: 是否粗体
            italic: 是否斜体
            font_name: 字体名称
            font_size: 字号（pt）

        Returns:
            如果找到匹配的 run 且格式匹配返回 True
        """
        from docx.shared import Pt

        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                if text in run.text:
                    if bold is not None and run.bold != bold:
                        continue
                    if italic is not None and run.italic != italic:
                        continue
                    if font_name is not None and run.font.name != font_name:
                        continue
                    if font_size is not None and run.font.size != Pt(font_size):
                        continue
                    return True
        return False


@dataclass
class DocumentFixture:
    """
    文档夹具，管理测试文档的生命周期

    Attributes:
        original_path: 原始夹具文件路径
        working_path: 工作副本路径（测试实际使用的文件）
        document_uri: Word Add-In 使用的文档 URI
        cleanup_on_success: 成功后是否清理
        document_closed: 文档是否已关闭（避免重复关闭）
    """

    original_path: Path
    working_path: Path
    document_uri: str
    cleanup_on_success: bool = True
    document_closed: bool = False

    def cleanup(self) -> None:
        """清理工作副本"""
        if self.working_path.exists():
            self.working_path.unlink()
            print(f"🧹 已清理: {self.working_path.name}")


@dataclass
class TestResult:
    """
    测试结果

    Attributes:
        success: 是否成功
        message: 结果消息
        data: 返回的数据
        error: 错误信息
        duration_ms: 执行时间（毫秒）
    """

    success: bool
    message: str
    data: dict[str, Any] | None = None
    error: str | None = None
    duration_ms: float = 0.0


@dataclass
class ExpectedStats:
    """
    预期的文档统计数据

    用于自动验证测试结果。
    设置为 None 的字段将跳过验证。
    """

    word_count: int | None = None
    character_count: int | None = None
    paragraph_count: int | None = None

    # 允许的误差范围（用于大文档等场景）
    word_count_tolerance: int = 0
    character_count_tolerance: int = 0
    paragraph_count_tolerance: int = 0


@dataclass
class ExpectedStructure:
    """
    预期的文档结构数据

    用于自动验证 word:get:documentStructure 测试结果。
    设置为 None 的字段将跳过验证。
    """

    paragraph_count: int | None = None
    table_count: int | None = None
    image_count: int | None = None
    section_count: int | None = None

    # 允许的误差范围
    paragraph_count_tolerance: int = 0
    table_count_tolerance: int = 0
    image_count_tolerance: int = 0
    section_count_tolerance: int = 0


# ==============================================================================
# Validator 类型定义
# ==============================================================================

# 传统验证器：仅接收协议返回的 data
DataValidator = Callable[[dict[str, Any]], bool]

# 内容验证器：接收 data 和 DocumentReader，用于双重验证
ContentValidator = Callable[[dict[str, Any], DocumentReader], bool]

# 统一验证器类型（向后兼容）
Validator = DataValidator | ContentValidator


def _call_validator(
    validator: Validator,
    data: dict[str, Any],
    reader: DocumentReader,
) -> bool:
    """
    智能调用验证器

    根据验证器的参数数量自动判断是传统模式还是双重验证模式。

    Args:
        validator: 验证函数
        data: 协议返回的数据
        reader: 文档读取器

    Returns:
        验证结果
    """
    sig = inspect.signature(validator)
    if len(sig.parameters) == 2:
        # 双重验证模式：传入 data 和 reader
        return validator(data, reader)  # type: ignore[call-arg]
    else:
        # 传统模式：仅传入 data
        return validator(data)  # type: ignore[call-arg]


@dataclass
class TestCase:
    """
    测试用例定义

    Attributes:
        name: 测试名称
        fixture_name: 夹具文件名（相对于 fixtures 目录）
        description: 测试描述
        expected: 预期结果
        validator: 自定义验证函数（可选）
            - DataValidator: (data) -> bool - 仅验证协议返回
            - ContentValidator: (data, reader) -> bool - 双重验证（协议 + 文档内容）
        expect_failure: 预期失败（哨兵测试）— 操作失败时判定为通过，
            若意外成功则提示能力可能已更新，需人工确认
    """

    name: str
    fixture_name: str
    description: str
    expected: ExpectedStats | None = None
    validator: Validator | None = None
    tags: list[str] = field(default_factory=list)
    expect_failure: bool = False


# ==============================================================================
# 文档操作工具
# ==============================================================================


def path_to_file_uri(path: Path) -> str:
    """
    将文件路径转换为 file:// URI

    Args:
        path: 文件路径

    Returns:
        file:// URI
    """
    # 确保路径是绝对路径
    abs_path = path.resolve()
    # URL 编码路径（保留 /）
    encoded = quote(str(abs_path), safe="/")
    return f"file://{encoded}"


def open_document(path: Path) -> subprocess.Popen[bytes] | None:
    """
    使用系统默认应用打开文档

    Args:
        path: 文档路径

    Returns:
        进程对象，如果失败则返回 None
    """
    system = platform.system()

    try:
        if system == "Darwin":  # macOS
            # 使用 open 命令，-a 指定应用（可选）
            process = subprocess.Popen(
                ["open", str(path)],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
        elif system == "Windows":
            # 使用 start 命令
            process = subprocess.Popen(
                ["cmd", "/c", "start", "", str(path)],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                shell=True,
            )
        else:  # Linux
            # 使用 xdg-open
            process = subprocess.Popen(
                ["xdg-open", str(path)],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
        return process
    except Exception as e:
        print(f"❌ 打开文档失败: {e}")
        return None


def save_document(path: Path) -> bool:
    """
    强制 Word 保存文档（通过 AppleScript）

    在用 python-docx 读取磁盘文件验证内容之前调用，
    确保 Word 内存中的修改已写入磁盘。

    Args:
        path: 文档路径

    Returns:
        是否成功
    """
    system = platform.system()

    if system != "Darwin":
        print("⚠️  自动保存仅支持 macOS")
        return False

    try:
        doc_name = path.name
        doc_stem = path.stem
        script = f'''
        tell application "Microsoft Word"
            -- 优先精确匹配
            set targetDocs to (every document whose name is "{doc_name}")
            if (count of targetDocs) is 0 then
                set targetDocs to (every document whose name starts with "{doc_stem}")
            end if
            repeat with d in targetDocs
                save d
            end repeat
            return count of targetDocs
        end tell
        '''
        result = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True,
            timeout=5,
        )
        if result.returncode == 0:
            saved_count = result.stdout.decode().strip()
            if saved_count != "0":
                return True
            else:
                print("⚠️  未找到匹配的文档进行保存")
                return False
        else:
            print(f"⚠️  AppleScript 保存错误: {result.stderr.decode()}")
            return False
    except Exception as e:
        print(f"⚠️  保存文档失败: {e}")
        return False


def close_document(path: Path) -> bool:
    """
    关闭文档（通过 AppleScript 或其他方式）

    注意：这个功能在不同平台上实现方式不同，
    且可能不完全可靠。

    Args:
        path: 文档路径

    Returns:
        是否成功
    """
    system = platform.system()

    try:
        if system == "Darwin":  # macOS
            # 使用文档名称匹配（而非路径，因为 Word 返回 HFS 格式路径）
            # 先尝试精确匹配完整文件名，失败则用 stem 做 contains 匹配
            # （兼容 Word 可能省略扩展名或追加 [Compatibility Mode] 等后缀的情况）
            doc_name = path.name
            doc_stem = path.stem
            script = f'''
            tell application "Microsoft Word"
                -- 优先精确匹配
                set targetDocs to (every document whose name is "{doc_name}")
                if (count of targetDocs) is 0 then
                    -- 降级: stem 前缀匹配 (无扩展名场景)
                    set targetDocs to (every document whose name starts with "{doc_stem}")
                end if
                repeat with d in targetDocs
                    close d saving no
                end repeat
                return count of targetDocs
            end tell
            '''
            result = subprocess.run(
                ["osascript", "-e", script],
                capture_output=True,
                timeout=5,
            )
            if result.returncode == 0:
                closed_count = result.stdout.decode().strip()
                print(f"📕 已关闭文档: {doc_name} ({closed_count} matched)")
                return True
            else:
                print(f"⚠️  AppleScript 返回错误: {result.stderr.decode()}")
                return False
        # 其他平台暂不支持自动关闭
        return False
    except Exception as e:
        print(f"⚠️  关闭文档失败: {e}")
        return False


# ==============================================================================
# Add-In 自动激活（macOS AppleScript）
# ==============================================================================


def _run_applescript(script: str, timeout: float = 15.0) -> tuple[bool, str]:
    """
    执行 AppleScript 并返回结果

    Args:
        script: AppleScript 脚本内容
        timeout: 超时时间（秒）

    Returns:
        (success, output): 是否成功和输出内容
    """
    try:
        result = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True,
            timeout=timeout,
        )
        output = result.stdout.decode().strip()
        if result.returncode == 0:
            return True, output
        else:
            stderr = result.stderr.decode().strip()
            return False, stderr
    except subprocess.TimeoutExpired:
        return False, "AppleScript 执行超时"
    except Exception as e:
        return False, str(e)


async def activate_word_addin(addin_name: str = DEFAULT_ADDIN_NAME) -> bool:
    """
    通过 AppleScript UI 自动化激活 Word Add-In（仅 macOS）

    自动完成两步操作：
    1. 点击 Word 功能区的「加载项」按钮
    2. 在弹出面板中点击目标 Add-In

    如果自动化失败，会打印日志提示用户手动操作。

    Args:
        addin_name: Add-In 名称（部分匹配即可）

    Returns:
        是否成功激活
    """
    if platform.system() != "Darwin":
        print("⚠️  自动激活 Add-In 仅支持 macOS，请手动激活")
        return False

    print(f"🔌 尝试自动激活 Add-In: {addin_name}...")

    # Phase 1: 点击「加载项」按钮
    # Word macOS UI 结构: front window → tab group 1 → scroll area → groups → AXButton "加载项"
    # 使用 entire contents 扁平化搜索，避免嵌套遍历失败
    # 功能区可能处于折叠状态（无 scroll area），需要先点击选项卡展开
    script_phase1 = '''
tell application "Microsoft Word" to activate
delay 0.5

tell application "System Events"
    tell process "Microsoft Word"
        set tg to tab group 1 of front window

        -- 检查功能区是否展开（scroll area 存在则展开）
        set ribbonExpanded to false
        try
            set saCount to count of scroll areas of tg
            if saCount > 0 then set ribbonExpanded to true
        end try

        -- 如果功能区折叠，点击「开始」展开
        if not ribbonExpanded then
            repeat with rb in radio buttons of tg
                if name of rb is "开始" or name of rb is "Home" then
                    click rb
                    delay 0.5
                    exit repeat
                end if
            end repeat
        end if

        -- 使用 entire contents 扁平搜索「加载项」按钮
        -- 使用 contains 匹配以兼容不同语言 (中文「加载项」/ English "Add-ins"/"Add-Ins")
        set allElems to entire contents of tg
        repeat with elem in allElems
            try
                if role of elem is "AXButton" then
                    set eName to name of elem
                    if eName contains "加载项" or eName contains "Add-in" then
                        click elem
                        return "ok"
                    end if
                end if
            end try
        end repeat

        return "not_found"
    end tell
end tell
'''

    ok, output = _run_applescript(script_phase1)
    if not ok or output != "ok":
        print(f"   ⚠️  未能自动点击「加载项」按钮 (原因: {output})")
        print(f"   👆 请手动点击「加载项」→「{addin_name}...」")
        return False

    print("   ✅ Phase 1: 已点击「加载项」按钮")

    # Phase 2: 提示用户手动点击目标 Add-In
    # Office 的加载项弹窗是自绘 UI（非原生 AX 元素），无法通过 AppleScript 自动化点击。
    # 如果文档已通过 Office.AutoShowTaskpaneWithDocument 预标记，则此弹窗不会出现，
    # Add-In 会自动加载，此函数返回 False 不影响后续流程。
    print(f"   👆 请在弹出的加载项面板中点击「{addin_name}...」")
    return False


def dump_word_ui_hierarchy() -> None:
    """
    调试工具：打印 Word 前端窗口的 UI 元素层级

    当自动激活失败时，运行此函数查看 Word 的 UI 结构，
    帮助调整 AppleScript 中的元素定位逻辑。

    Usage:
        python -c "from manual_tests.e2e_base import dump_word_ui_hierarchy; dump_word_ui_hierarchy()"
    """
    if platform.system() != "Darwin":
        print("仅支持 macOS")
        return

    script = '''
tell application "Microsoft Word" to activate
delay 0.5

tell application "System Events"
    tell process "Microsoft Word"
        set output to ""
        set uiElems to UI elements of front window
        repeat with elem in uiElems
            try
                set elemRole to role of elem
                set elemName to name of elem
                set elemDesc to description of elem
                set output to output & "[" & elemRole & "] name=" & elemName & " desc=" & elemDesc & linefeed
                -- 一层子元素
                try
                    set children to UI elements of elem
                    repeat with child in children
                        try
                            set cRole to role of child
                            set cName to name of child
                            set cDesc to description of child
                            set output to output & "  [" & cRole & "] name=" & cName & " desc=" & cDesc & linefeed
                        end try
                    end repeat
                end try
            end try
        end repeat
        return output
    end tell
end tell
'''

    ok, output = _run_applescript(script, timeout=30.0)
    if ok:
        print("=== Word UI Hierarchy ===")
        print(output)
    else:
        print(f"获取 UI 层级失败: {output}")


# ==============================================================================
# E2E 测试运行器
# ==============================================================================


class E2ETestRunner:
    """
    E2E 测试运行器

    提供自动化的文档夹具管理、Workspace 连接、测试执行和清理功能。
    """

    def __init__(
        self,
        fixtures_dir: str | Path | None = None,
        host: str = "127.0.0.1",
        port: int = 3000,
        connection_timeout: float = 30.0,
        auto_open: bool = True,
        auto_close: bool = True,
        auto_activate: bool = True,
        cleanup_on_success: bool = True,
    ):
        """
        初始化测试运行器

        Args:
            fixtures_dir: 夹具目录（默认使用 FIXTURES_ROOT）
            host: Workspace 服务器地址
            port: Workspace 服务器端口
            connection_timeout: Add-In 连接超时时间（秒）
            auto_open: 是否自动打开文档
            auto_close: 是否自动关闭文档（独立于 auto_open，调试时可只禁用关闭）
            auto_activate: 是否自动激活 Add-In（通过 AppleScript，仅 macOS）
            cleanup_on_success: 成功后是否清理测试副本
        """
        self.fixtures_dir = Path(fixtures_dir) if fixtures_dir else FIXTURES_ROOT
        self.host = host
        self.port = port
        self.connection_timeout = connection_timeout
        self.auto_open = auto_open
        self.auto_close = auto_close
        self.auto_activate = auto_activate
        self.cleanup_on_success = cleanup_on_success

        # 确保临时目录存在
        TEMP_ROOT.mkdir(parents=True, exist_ok=True)

    def _create_working_copy(self, fixture_path: Path) -> Path:
        """
        创建夹具文件的工作副本

        Args:
            fixture_path: 原始夹具文件路径

        Returns:
            工作副本路径
        """
        # 生成唯一的文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        working_name = f"{fixture_path.stem}_{timestamp}{fixture_path.suffix}"
        working_path = TEMP_ROOT / working_name

        # 复制文件
        shutil.copy2(fixture_path, working_path)
        return working_path

    @asynccontextmanager
    async def prepare_document(
        self,
        fixture_name: str,
        open_delay: float = 2.0,
    ) -> AsyncIterator[DocumentFixture]:
        """
        准备测试文档（复制、打开）

        Args:
            fixture_name: 夹具文件名（相对于 fixtures_dir）
            open_delay: 打开后等待时间（秒），让 Office 加载插件

        Yields:
            DocumentFixture: 文档夹具对象

        Example:
            async with runner.prepare_document("empty.docx") as fixture:
                print(fixture.document_uri)
        """
        # 解析夹具路径
        fixture_path = self.fixtures_dir / fixture_name
        if not fixture_path.exists():
            raise FileNotFoundError(f"夹具文件不存在: {fixture_path}")

        # 创建工作副本
        working_path = self._create_working_copy(fixture_path)
        print(f"📄 创建工作副本: {working_path.name}")

        # 生成 document URI
        document_uri = path_to_file_uri(working_path)

        fixture = DocumentFixture(
            original_path=fixture_path,
            working_path=working_path,
            document_uri=document_uri,
            cleanup_on_success=self.cleanup_on_success,
        )

        success = False
        try:
            # 自动打开文档
            if self.auto_open:
                print(f"📂 打开文档: {working_path.name}")
                open_document(working_path)
                await asyncio.sleep(open_delay)

            yield fixture
            success = True

        finally:
            # 清理逻辑
            if success and fixture.cleanup_on_success:
                # 尝试关闭文档（如果还没关闭的话）
                if self.auto_close and not fixture.document_closed:
                    close_document(working_path)
                    fixture.document_closed = True
                    await asyncio.sleep(0.5)
                fixture.cleanup()
            elif not success:
                print(f"⚠️  测试失败，保留文件供调试: {working_path}")

    @asynccontextmanager
    async def run_with_workspace(
        self,
        fixture_name: str,
        open_delay: float = 2.0,
    ) -> AsyncIterator[tuple[OfficeWorkspace, DocumentFixture]]:
        """
        准备文档并启动 Workspace

        这是最常用的上下文管理器，一次性完成：
        1. 复制测试文档
        2. 打开文档
        3. 启动 Workspace
        4. 等待 Add-In 连接

        Args:
            fixture_name: 夹具文件名
            open_delay: 打开后等待时间

        Yields:
            (workspace, fixture): Workspace 实例和文档夹具

        Example:
            async with runner.run_with_workspace("simple.docx") as (workspace, fixture):
                result = await workspace.execute(action)
        """
        async with self.prepare_document(fixture_name, open_delay) as fixture:  # pyright: ignore[reportGeneralTypeIssues]
            workspace = OfficeWorkspace(host=self.host, port=self.port)
            try:
                await workspace.start()
                print("✅ Workspace 启动成功")

                # 自动激活 Add-In（仅 macOS，需要辅助功能权限）
                if self.auto_open and self.auto_activate:
                    activated = await activate_word_addin()
                    if not activated:
                        print("👆 请手动点击「加载项」→「word-editor...」激活 Add-In")

                # 等待 Add-In 连接
                print("⏳ 等待 Word Add-In 连接...")
                connected = await workspace.wait_for_addin_connection(
                    timeout=self.connection_timeout
                )
                if not connected:
                    raise RuntimeError("超时：未检测到 Add-In 连接")
                print("✅ Add-In 已连接")

                yield workspace, fixture

            finally:
                # 先关闭文档，让 Add-In 断开连接
                # 这样 workspace.stop() 就不用等待连接超时
                # 如果 cleanup_on_success=False，保留文档供人工检查
                if self.auto_close and not fixture.document_closed and fixture.cleanup_on_success:
                    close_document(fixture.working_path)
                    fixture.document_closed = True
                    await asyncio.sleep(0.5)  # 等待 Add-In 断开
                await workspace.stop()

    def verify_stats(
        self,
        actual: dict[str, Any],
        expected: ExpectedStats,
    ) -> tuple[bool, list[str]]:
        """
        验证文档统计数据

        Args:
            actual: 实际返回的统计数据
            expected: 预期的统计数据

        Returns:
            (success, messages): 是否通过和验证消息列表
        """
        messages: list[str] = []
        success = True

        # 验证字数
        if expected.word_count is not None:
            actual_wc = actual.get("wordCount", 0)
            diff = abs(actual_wc - expected.word_count)
            if diff <= expected.word_count_tolerance:
                messages.append(f"✅ 字数: {actual_wc} (预期 {expected.word_count})")
            else:
                messages.append(
                    f"❌ 字数不匹配: {actual_wc} (预期 {expected.word_count}, "
                    f"误差 {diff} > 允许 {expected.word_count_tolerance})"
                )
                success = False

        # 验证字符数
        if expected.character_count is not None:
            actual_cc = actual.get("characterCount", 0)
            diff = abs(actual_cc - expected.character_count)
            if diff <= expected.character_count_tolerance:
                messages.append(f"✅ 字符数: {actual_cc} (预期 {expected.character_count})")
            else:
                messages.append(
                    f"❌ 字符数不匹配: {actual_cc} (预期 {expected.character_count}, "
                    f"误差 {diff} > 允许 {expected.character_count_tolerance})"
                )
                success = False

        # 验证段落数
        if expected.paragraph_count is not None:
            actual_pc = actual.get("paragraphCount", 0)
            diff = abs(actual_pc - expected.paragraph_count)
            if diff <= expected.paragraph_count_tolerance:
                messages.append(f"✅ 段落数: {actual_pc} (预期 {expected.paragraph_count})")
            else:
                messages.append(
                    f"❌ 段落数不匹配: {actual_pc} (预期 {expected.paragraph_count}, "
                    f"误差 {diff} > 允许 {expected.paragraph_count_tolerance})"
                )
                success = False

        return success, messages

    def verify_structure(
        self,
        actual: dict[str, Any],
        expected: "ExpectedStructure",
    ) -> tuple[bool, list[str]]:
        """
        验证文档结构数据

        Args:
            actual: 实际返回的结构数据
            expected: 预期的结构数据

        Returns:
            (success, messages): 是否通过和验证消息列表
        """
        messages: list[str] = []
        success = True

        # 验证段落数
        if expected.paragraph_count is not None:
            actual_pc = actual.get("paragraphCount", 0)
            diff = abs(actual_pc - expected.paragraph_count)
            if diff <= expected.paragraph_count_tolerance:
                messages.append(f"✅ 段落数: {actual_pc} (预期 {expected.paragraph_count})")
            else:
                messages.append(
                    f"❌ 段落数不匹配: {actual_pc} (预期 {expected.paragraph_count}, "
                    f"误差 {diff} > 允许 {expected.paragraph_count_tolerance})"
                )
                success = False

        # 验证表格数
        if expected.table_count is not None:
            actual_tc = actual.get("tableCount", 0)
            diff = abs(actual_tc - expected.table_count)
            if diff <= expected.table_count_tolerance:
                messages.append(f"✅ 表格数: {actual_tc} (预期 {expected.table_count})")
            else:
                messages.append(
                    f"❌ 表格数不匹配: {actual_tc} (预期 {expected.table_count}, "
                    f"误差 {diff} > 允许 {expected.table_count_tolerance})"
                )
                success = False

        # 验证图片数
        if expected.image_count is not None:
            actual_ic = actual.get("imageCount", 0)
            diff = abs(actual_ic - expected.image_count)
            if diff <= expected.image_count_tolerance:
                messages.append(f"✅ 图片数: {actual_ic} (预期 {expected.image_count})")
            else:
                messages.append(
                    f"❌ 图片数不匹配: {actual_ic} (预期 {expected.image_count}, "
                    f"误差 {diff} > 允许 {expected.image_count_tolerance})"
                )
                success = False

        # 验证节数
        if expected.section_count is not None:
            actual_sc = actual.get("sectionCount", 0)
            diff = abs(actual_sc - expected.section_count)
            if diff <= expected.section_count_tolerance:
                messages.append(f"✅ 节数: {actual_sc} (预期 {expected.section_count})")
            else:
                messages.append(
                    f"❌ 节数不匹配: {actual_sc} (预期 {expected.section_count}, "
                    f"误差 {diff} > 允许 {expected.section_count_tolerance})"
                )
                success = False

        return success, messages


# ==============================================================================
# 测试夹具创建工具
# ==============================================================================


def create_empty_docx(path: Path) -> None:
    """创建空白 Word 文档"""
    from docx import Document

    doc = Document()
    doc.save(str(path))


def create_simple_docx(path: Path) -> None:
    """
    创建简单文本文档

    内容：
    - 标题
    - 3 个段落，每段若干句话
    """
    from docx import Document

    doc = Document()

    # 标题
    doc.add_heading("测试文档", level=1)

    # 段落
    paragraphs = [
        "这是第一个段落。它包含一些简单的中文文本。这段文字用于测试文档统计功能。",
        "第二个段落在这里。我们需要确保字数统计是准确的。Word 的统计功能会计算所有文字。",
        "最后一个段落。测试结束。谢谢使用 Office4AI。",
    ]

    for para in paragraphs:
        doc.add_paragraph(para)

    doc.save(str(path))


def create_complex_docx(path: Path) -> None:
    """
    创建复杂文档（包含表格、列表等）

    内容：
    - 标题
    - 段落
    - 表格
    - 列表
    """
    from docx import Document

    doc = Document()

    # 标题
    doc.add_heading("复杂测试文档", level=1)

    # 段落
    doc.add_paragraph("这是一个包含多种元素的复杂文档。")

    # 子标题
    doc.add_heading("表格示例", level=2)

    # 表格
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"单元格 {i+1}-{j+1}"

    # 列表
    doc.add_heading("列表示例", level=2)
    items = ["第一项", "第二项", "第三项"]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    # 结尾段落
    doc.add_paragraph("这是文档的结尾。")

    doc.save(str(path))


def create_large_docx(path: Path, pages: int = 10) -> None:
    """
    创建大文档（多页）

    Args:
        path: 保存路径
        pages: 页数（约估）
    """
    from docx import Document

    doc = Document()

    doc.add_heading("大型测试文档", level=1)

    # 每页大约需要 500-600 字
    para_text = (
        "这是一段用于填充大型文档的文本。"
        "我们需要足够多的内容来测试文档统计功能的性能。"
        "Word 文档可能包含大量的文字、段落和其他元素。"
        "这个测试旨在验证系统在处理大型文档时的表现。"
    ) * 5

    # 生成足够的段落
    for i in range(pages * 3):  # 每页约 3 段
        doc.add_paragraph(f"第 {i+1} 段：{para_text}")

    doc.save(str(path))


def ensure_fixtures(fixture_dir: Path) -> dict[str, Path]:
    """
    确保所有测试夹具文件存在

    如果文件不存在，则创建它们。

    Args:
        fixture_dir: 夹具目录

    Returns:
        夹具文件路径字典
    """
    fixture_dir.mkdir(parents=True, exist_ok=True)

    fixtures = {
        "empty.docx": create_empty_docx,
        "simple.docx": create_simple_docx,
        "complex.docx": create_complex_docx,
        "large.docx": lambda p: create_large_docx(p, pages=10),
    }

    paths: dict[str, Path] = {}
    for name, creator in fixtures.items():
        path = fixture_dir / name
        if not path.exists():
            print(f"📝 创建夹具: {name}")
            creator(path)
        paths[name] = path

    return paths


# ==============================================================================
# 命令行工具
# ==============================================================================


def main() -> None:
    """命令行入口：创建测试夹具"""
    import argparse

    parser = argparse.ArgumentParser(description="E2E 测试夹具管理工具")
    parser.add_argument(
        "--create-fixtures",
        metavar="DIR",
        help="创建测试夹具到指定目录",
    )
    parser.add_argument(
        "--clean-temp",
        action="store_true",
        help="清理临时测试文件",
    )

    args = parser.parse_args()

    if args.create_fixtures:
        fixture_dir = Path(args.create_fixtures)
        paths = ensure_fixtures(fixture_dir)
        print(f"\n✅ 已创建 {len(paths)} 个夹具文件到: {fixture_dir}")
        for name, path in paths.items():
            print(f"   - {name}: {path}")

    if args.clean_temp:
        if TEMP_ROOT.exists():
            count = len(list(TEMP_ROOT.glob("*.docx")))
            shutil.rmtree(TEMP_ROOT)
            print(f"🧹 已清理 {count} 个临时文件")
        else:
            print("📭 无临时文件需要清理")


if __name__ == "__main__":
    main()
